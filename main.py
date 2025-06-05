from docx import Document
from utils import iter_block_items,extract_styles, align_blocks, identify_changes, split_and_style
from itertools import zip_longest
import difflib


# name, attribute
ATTRIBUTES_TO_CHECK = [
    ("bold", "bold"),
    ("italic", "italic"),
    ("underline", "underline"),
    ("font color", "font.color.rgb"),
    ("font size", "font.size")
]


def compare_documents(old_document_path: str, new_document_path: str):
    print("⦿⦿⦿⦿⦿⦿⦿ Refreshing ⦿⦿⦿⦿⦿⦿⦿\n")
    # take in two files and use difflib to detect differences
    old_document = Document(old_document_path)
    new_document = Document(new_document_path)
    old_content = [para.text for para in old_document.paragraphs if para.text != ""]
    new_content = [para.text for para in new_document.paragraphs if para.text != ""]
    diff = difflib.unified_diff(old_content, new_content)
    
    changes = []
    removed = []
    # check for the changes in text only
    for line in diff:
        filtered_line = line[1:].lstrip('+-').rstrip('.')
        if line.startswith('-') and not line.startswith('---'):
            removed.append(filtered_line)
            changes.append(f"🔴 Removed: {filtered_line}")
        elif line.startswith('+') and not line.startswith('+++'):
            found = False
            for removed_line in removed:
                if removed_line in filtered_line or filtered_line in removed_line:
                    found = True
                    removed.remove(removed_line)
                    changes.append(f"🔵 From: {removed_line}\n🔵 Updated to: {filtered_line}")
                    changes.remove(f"🔴 Removed: {removed_line}")
            if not found:
                changes.append(f"🟢 Added: {filtered_line}")

    # check for changes in stylings (bold, italic, underscore, color, size)
    old_block_items = list(iter_block_items(old_document))
    new_block_items = list(iter_block_items(new_document))
    aligned_pairs = align_blocks(old_block_items, new_block_items)

    block_num = 0
    for idx, (old_block, new_block) in enumerate(aligned_pairs):
        block_num += 1
        old_styling = extract_styles(old_block)
        new_styling = extract_styles(new_block)   
        old_words_styling = []
        new_words_styling = []

        for (old_text, old_style), (new_text, new_style) in zip_longest(old_styling, new_styling, fillvalue=(None, None)):
            old_words_styling.extend(split_and_style(old_text, old_style))
            new_words_styling.extend(split_and_style(new_text, new_style))

        matched, changed, added, deleted = identify_changes(old_words_styling, new_words_styling)
        
        if changed: 
            temp_change = {}
            temp_add = {}
            temp_remove = {}
            # go through all the changes and check for the changes in attributes
            change_num = 0
            for change in changed:
                change_num += 1
                new_attrs = change[1][1]
                old_attrs = change[0][1]
                attr_name = change[0][0]

                for attr in new_attrs:
                    if attr in old_attrs:
                        new_attr = new_attrs[attr]

                        # check for stylings that have been changed
                        if old_attrs[attr] != new_attrs[attr]:
                            if (attr, new_attrs[attr]) in temp_change:
                                temp_change[(attr, new_attr)].append((attr_name, old_attrs[attr]))
                            else:
                                temp_change[(attr, new_attrs[attr])] = [(attr_name, old_attrs[attr])]
                    else:
                        if (attr, new_attrs[attr]) in temp_add:
                                temp_add[(attr, new_attrs[attr])].append((attr_name))
                        else:
                            temp_add[(attr, new_attrs[attr])] = [(attr_name)]

                # check for stylings that have been removed
                for attr in old_attrs:
                    if attr not in new_attrs:
                        # changes.append(f"🟡 Removed {change_num} [{attr}]: '{change[0][0]}'")
                        if attr in temp_remove:
                            temp_remove[attr].append(change[0][0])
                        else:
                            temp_remove[attr] = [change[0][0]]
            
            # update the attribute removal
            for attr, change in temp_remove.items():
                changes.append(f"🟡 Removed [{attr}]: '{" ".join([x for x in change])}'")
            
            # update the attribute changes
            for attr, change in temp_change.items():
                new_attr = attr[0]
                old_attr = attr[1]
                attr_name = change[0][1]
                changes.append(f"🟡 Changed [{old_attr}]: '{" ".join([x[0] for x in change])}': {attr_name} to {new_attr}")

            # update the new attributes
            for attr, add in temp_add.items():
                new_attr = attr[0]
                old_attr = attr[1]
                attr_name = change[0][1]
                if old_attr is True:
                    changes.append(f"🟡 Added [{new_attr}]: '{" ".join([x for x in add])}'")
                else:
                    changes.append(f"🟡 Added [{new_attr}]: '{" ".join([x for x in add])}': {old_attr}")

    for change in changes:
        print(change, "\n")

    return changes

if __name__ == "__main__":
    compare_documents("documents/file1.docx", "documents/file2.docx")